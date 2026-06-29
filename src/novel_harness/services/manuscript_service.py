"""Volume/chapter organization and deterministic manuscript export."""

from __future__ import annotations

import io
import json
import re
import zipfile
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from novel_harness.models import (
    ManuscriptChapter,
    ManuscriptOutline,
    ManuscriptPreview,
    ManuscriptVolume,
)
from novel_harness.storage.repositories import Repositories


class ManuscriptService:
    def __init__(self, session: Session) -> None:
        self.session = session
        self.repositories = Repositories(session)

    def outline(
        self,
        project_id: str,
        *,
        include_archived: bool = False,
    ) -> ManuscriptOutline:
        self.repositories.projects.require(project_id)
        self.ensure_default_volume(project_id)
        volumes = self.repositories.manuscript_volumes.list_ordered(
            project_id,
            include_archived=include_archived,
        )
        visible_ids = {volume.id for volume in volumes}
        chapters = [
            chapter
            for chapter in self.repositories.manuscript_chapters.list_ordered(project_id)
            if chapter.volume_id in visible_ids
        ]
        volume_positions = {volume.id: index for index, volume in enumerate(volumes)}
        chapters.sort(
            key=lambda chapter: (
                volume_positions[chapter.volume_id],
                chapter.position,
                chapter.created_at,
            )
        )
        return ManuscriptOutline(volumes=volumes, chapters=chapters)

    def ensure_default_volume(self, project_id: str) -> ManuscriptVolume:
        volumes = self.repositories.manuscript_volumes.list_ordered(
            project_id,
            include_archived=True,
        )
        if volumes:
            return volumes[0]
        return self.repositories.manuscript_volumes.add(
            ManuscriptVolume(
                project_id=project_id,
                title="第一卷",
                position=1,
            )
        )

    def create_volume(
        self,
        project_id: str,
        *,
        title: str,
        description: str = "",
        position: int | None = None,
    ) -> ManuscriptVolume:
        self.repositories.projects.require(project_id)
        volume = ManuscriptVolume(
            project_id=project_id,
            title=title,
            description=description,
            position=position or self.repositories.manuscript_volumes.next_position(project_id),
        )
        return self.repositories.manuscript_volumes.add(volume)

    def update_volume(self, volume_id: str, **changes: Any) -> ManuscriptVolume:
        volume = self.repositories.manuscript_volumes.require(volume_id)
        self.repositories.projects.require(volume.project_id)
        allowed = {"title", "description", "position", "status"}
        self._validate_changes(changes, allowed)
        if changes.get("status") == "archived":
            active = self.repositories.manuscript_volumes.list_ordered(volume.project_id)
            if len(active) <= 1:
                raise ValueError("作品至少需要保留一个有效卷")
        values = {**volume.model_dump(), **changes}
        return self.repositories.manuscript_volumes.update(ManuscriptVolume.model_validate(values))

    def reorder_volumes(
        self,
        project_id: str,
        ordered_ids: list[str],
    ) -> list[ManuscriptVolume]:
        self.repositories.projects.require(project_id)
        volumes = self.repositories.manuscript_volumes.list_ordered(project_id)
        self._require_same_ids(ordered_ids, [item.id for item in volumes], "卷")
        by_id = {item.id: item for item in volumes}
        return [
            self.repositories.manuscript_volumes.update(
                ManuscriptVolume.model_validate(
                    {**by_id[item_id].model_dump(), "position": position}
                )
            )
            for position, item_id in enumerate(ordered_ids, 1)
        ]

    def create_chapter(
        self,
        project_id: str,
        *,
        volume_id: str,
        title: str,
        summary: str = "",
        position: int | None = None,
        draft_id: str | None = None,
        status: str | None = None,
    ) -> ManuscriptChapter:
        self.repositories.projects.require(project_id)
        volume = self._require_active_volume(volume_id, project_id)
        self._ensure_draft_available(draft_id)
        draft_status = self._draft_status(draft_id, project_id)
        chapter_status = status or self._status_for_draft(draft_status)
        accepted_draft_id = draft_id if draft_status == "accepted" else None
        self._validate_chapter_status(chapter_status, accepted_draft_id)
        chapter = ManuscriptChapter(
            project_id=project_id,
            volume_id=volume.id,
            title=title,
            summary=summary,
            position=position
            or self.repositories.manuscript_chapters.next_position(project_id, volume.id),
            draft_id=draft_id,
            accepted_draft_id=accepted_draft_id,
            status=chapter_status,
        )
        return self.repositories.manuscript_chapters.add(chapter)

    def update_chapter(self, chapter_id: str, **changes: Any) -> ManuscriptChapter:
        chapter = self.repositories.manuscript_chapters.require(chapter_id)
        self.repositories.projects.require(chapter.project_id)
        allowed = {"volume_id", "title", "summary", "position", "draft_id", "status"}
        self._validate_changes(changes, allowed)
        volume_id = str(changes.get("volume_id") or chapter.volume_id)
        self._require_active_volume(volume_id, chapter.project_id)
        draft_id = changes["draft_id"] if "draft_id" in changes else chapter.draft_id
        self._ensure_draft_available(draft_id, chapter_id=chapter.id)
        draft_status = self._draft_status(draft_id, chapter.project_id)
        accepted_draft_id = chapter.accepted_draft_id
        if draft_status == "accepted":
            accepted_draft_id = draft_id
        if "draft_id" in changes and "status" not in changes:
            changes["status"] = self._status_for_draft(draft_status)
        chapter_status = str(changes.get("status") or chapter.status)
        self._validate_chapter_status(chapter_status, accepted_draft_id)
        values = {
            **chapter.model_dump(),
            **changes,
            "volume_id": volume_id,
            "accepted_draft_id": accepted_draft_id,
        }
        return self.repositories.manuscript_chapters.update(
            ManuscriptChapter.model_validate(values)
        )

    def require_chapter(self, project_id: str, chapter_id: str) -> ManuscriptChapter:
        self.repositories.projects.require(project_id)
        chapter = self.repositories.manuscript_chapters.require(chapter_id)
        if chapter.project_id != project_id:
            raise ValueError("章节不属于当前作品")
        self._require_active_volume(chapter.volume_id, project_id)
        return chapter

    def attach_draft(self, chapter_id: str, draft_id: str) -> ManuscriptChapter:
        chapter = self.repositories.manuscript_chapters.require(chapter_id)
        draft_status = self._draft_status(draft_id, chapter.project_id)
        self._ensure_draft_available(draft_id, chapter_id=chapter.id)
        accepted_draft_id = draft_id if draft_status == "accepted" else chapter.accepted_draft_id
        return self.repositories.manuscript_chapters.update(
            chapter.model_copy(
                update={
                    "draft_id": draft_id,
                    "accepted_draft_id": accepted_draft_id,
                    "status": self._status_for_draft(draft_status),
                }
            )
        )

    def reject_current_draft(self, draft_id: str) -> ManuscriptChapter | None:
        chapter = self.repositories.manuscript_chapters.get_by_draft(draft_id)
        if chapter is None or chapter.draft_id != draft_id:
            return chapter
        fallback = chapter.accepted_draft_id
        return self.repositories.manuscript_chapters.update(
            chapter.model_copy(
                update={
                    "draft_id": fallback,
                    "status": "accepted" if fallback else "planned",
                }
            )
        )

    def reorder_chapters(
        self,
        project_id: str,
        volume_id: str,
        ordered_ids: list[str],
    ) -> list[ManuscriptChapter]:
        self.repositories.projects.require(project_id)
        self._require_active_volume(volume_id, project_id)
        chapters = self.repositories.manuscript_chapters.list_ordered(
            project_id,
            volume_id=volume_id,
        )
        self._require_same_ids(ordered_ids, [item.id for item in chapters], "章节")
        by_id = {item.id: item for item in chapters}
        return [
            self.repositories.manuscript_chapters.update(
                ManuscriptChapter.model_validate(
                    {**by_id[item_id].model_dump(), "position": position}
                )
            )
            for position, item_id in enumerate(ordered_ids, 1)
        ]

    def export_markdown(self, project_id: str, object_store: Any) -> tuple[str, int]:
        project = self.repositories.projects.require(project_id)
        outline = self.outline(project_id)
        sections = [f"# {project.name}", ""]
        if project.premise:
            sections.extend([project.premise.strip(), ""])
        count = 0
        for volume in outline.volumes:
            volume_chapters = [
                chapter
                for chapter in outline.chapters
                if chapter.volume_id == volume.id and self._export_draft_id(chapter)
            ]
            if not volume_chapters:
                continue
            sections.extend([f"## {volume.title}", ""])
            if volume.description:
                sections.extend([volume.description.strip(), ""])
            for chapter in volume_chapters:
                body = self._accepted_body(chapter, object_store)
                sections.extend([f"### {chapter.title}", "", body.strip(), ""])
                count += 1
        if count == 0:
            raise ValueError("作品中还没有可导出的已接受章节")
        return "\n".join(sections).rstrip() + "\n", count

    def export_zip(self, project_id: str, object_store: Any) -> tuple[bytes, int]:
        project = self.repositories.projects.require(project_id)
        outline = self.outline(project_id)
        manuscript, count = self.export_markdown(project_id, object_store)
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("manuscript.md", manuscript)
            archive.writestr(
                "metadata.json",
                json.dumps(
                    {
                        "project": project.model_dump(mode="json"),
                        "volumes": [item.model_dump(mode="json") for item in outline.volumes],
                        "chapters": [item.model_dump(mode="json") for item in outline.chapters],
                        "exported_chapters": count,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            )
            volume_numbers = {volume.id: index for index, volume in enumerate(outline.volumes, 1)}
            for chapter in outline.chapters:
                if not self._export_draft_id(chapter) or chapter.volume_id not in volume_numbers:
                    continue
                body = self._accepted_body(chapter, object_store)
                volume_number = volume_numbers[chapter.volume_id]
                path = (
                    f"chapters/{volume_number:02d}/"
                    f"{chapter.position:03d}-{self._safe_name(chapter.title)}.md"
                )
                archive.writestr(path, f"# {chapter.title}\n\n{body.strip()}\n")
        return buffer.getvalue(), count

    def export_docx(self, project_id: str, object_store: Any) -> tuple[bytes, int]:
        project = self.repositories.projects.require(project_id)
        outline = self.outline(project_id)
        document = Document()
        document.add_heading(project.name, level=0)
        if project.premise:
            document.add_paragraph(project.premise.strip())
        count = 0
        for volume in outline.volumes:
            chapters = [
                chapter
                for chapter in outline.chapters
                if chapter.volume_id == volume.id and self._export_draft_id(chapter)
            ]
            if not chapters:
                continue
            document.add_heading(volume.title, level=1)
            if volume.description:
                document.add_paragraph(volume.description.strip())
            for chapter in chapters:
                document.add_heading(chapter.title, level=2)
                body = self._accepted_body(chapter, object_store)
                for paragraph in self._paragraphs(body):
                    document.add_paragraph(paragraph)
                count += 1
        if count == 0:
            raise ValueError("作品中还没有可导出的已接受章节")
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue(), count

    def preview(self, project_id: str, object_store: Any) -> ManuscriptPreview:
        outline = self.outline(project_id)
        total_characters = 0
        total_paragraphs = 0
        exportable = 0
        for chapter in outline.chapters:
            if not self._export_draft_id(chapter):
                continue
            body = self._accepted_body(chapter, object_store)
            total_characters += len("".join(body.split()))
            total_paragraphs += len(self._paragraphs(body))
            exportable += 1
        return ManuscriptPreview(
            volume_count=len(outline.volumes),
            chapter_count=len(outline.chapters),
            exportable_chapter_count=exportable,
            total_characters=total_characters,
            total_paragraphs=total_paragraphs,
        )

    def _require_active_volume(self, volume_id: str, project_id: str) -> ManuscriptVolume:
        volume = self.repositories.manuscript_volumes.require(volume_id)
        if volume.project_id != project_id:
            raise ValueError("卷不属于当前作品")
        if volume.status != "active":
            raise ValueError("归档卷不能新增或移动章节")
        return volume

    def _draft_status(self, draft_id: str | None, project_id: str) -> str | None:
        if draft_id is None:
            return None
        draft = self.repositories.generations.require(draft_id)
        if draft.project_id != project_id:
            raise ValueError("草稿不属于当前作品")
        return draft.status

    def _ensure_draft_available(
        self,
        draft_id: str | None,
        *,
        chapter_id: str | None = None,
    ) -> None:
        if draft_id is None:
            return
        linked = self.repositories.manuscript_chapters.get_by_draft(draft_id)
        if linked is not None and linked.id != chapter_id:
            raise ValueError("该草稿已经关联到其他章节")

    def _accepted_body(self, chapter: ManuscriptChapter, object_store: Any) -> str:
        draft_id = self._export_draft_id(chapter)
        if draft_id is None:
            raise ValueError(f"章节《{chapter.title}》没有已接受版本")
        draft = self.repositories.generations.require(draft_id)
        if draft.status != "accepted":
            raise ValueError(f"章节《{chapter.title}》关联的草稿尚未接受")
        if not draft.object_key:
            raise ValueError(f"章节《{chapter.title}》缺少正文文件")
        return object_store.get_bytes(draft.object_key).decode("utf-8")

    @staticmethod
    def _status_for_draft(draft_status: str | None) -> str:
        if draft_status is None:
            return "planned"
        return "accepted" if draft_status == "accepted" else "drafting"

    @staticmethod
    def _validate_chapter_status(status: str, accepted_draft_id: str | None) -> None:
        if status in {"accepted", "completed"} and accepted_draft_id is None:
            raise ValueError("章节标记为已接受或已完成前，必须关联已接受草稿")

    @staticmethod
    def _validate_changes(changes: dict[str, Any], allowed: set[str]) -> None:
        unknown = set(changes) - allowed
        if unknown:
            raise ValueError(f"unsupported manuscript fields: {sorted(unknown)}")

    @staticmethod
    def _require_same_ids(actual: list[str], expected: list[str], label: str) -> None:
        if len(actual) != len(set(actual)) or set(actual) != set(expected):
            raise ValueError(f"{label}排序必须包含当前列表中的全部且不重复的 ID")

    @staticmethod
    def _safe_name(value: str) -> str:
        name = re.sub(r'[<>:"/\\|?*\\x00-\\x1f]', "-", value).strip(" .")
        return name[:80] or "chapter"

    @staticmethod
    def _export_draft_id(chapter: ManuscriptChapter) -> str | None:
        if chapter.accepted_draft_id:
            return chapter.accepted_draft_id
        if chapter.status in {"accepted", "completed"}:
            return chapter.draft_id
        return None

    @staticmethod
    def _paragraphs(body: str) -> list[str]:
        return [item.strip() for item in re.split(r"\n\s*\n", body) if item.strip()]
