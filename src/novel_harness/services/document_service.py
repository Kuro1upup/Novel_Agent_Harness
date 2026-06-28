"""Safe document ingestion with MinIO/Milvus compensation."""

from __future__ import annotations

import hashlib
import io
import mimetypes
import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from novel_harness.exceptions import DocumentError
from novel_harness.models import Document, DocumentChunk
from novel_harness.providers.vectorstore import VectorRecord
from novel_harness.storage.repositories import Repositories


class DocumentService:
    SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".docx", ".pdf"}

    def __init__(
        self,
        session: Session,
        *,
        object_store: Any,
        vector_store: Any | None = None,
        embedding_provider: Any | None = None,
        max_upload_bytes: int = 20 * 1024 * 1024,
    ) -> None:
        self.repositories = Repositories(session)
        self.object_store = object_store
        self.vector_store = vector_store
        self.embedding_provider = embedding_provider
        self.max_upload_bytes = max_upload_bytes

    def ingest_path(self, project_id: str, path: Path) -> tuple[Document, str]:
        if not path.is_file():
            raise DocumentError(f"file does not exist: {path}")
        return self.ingest_bytes(project_id, path.name, path.read_bytes())

    def ingest_bytes(
        self,
        project_id: str,
        filename: str,
        content: bytes | str,
        *,
        mime_type: str | None = None,
    ) -> tuple[Document, str]:
        self.repositories.projects.require(project_id)
        if isinstance(content, str):
            content = content.encode("utf-8")
        safe_name = self._safe_filename(filename)
        suffix = Path(safe_name).suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise DocumentError(f"unsupported document type: {suffix or '<none>'}")
        if not content:
            raise DocumentError("document is empty")
        if len(content) > self.max_upload_bytes:
            raise DocumentError("document exceeds configured upload limit")
        digest = hashlib.sha256(content).hexdigest()
        existing = self.repositories.documents.get_by_hash(project_id, digest)
        if existing is not None:
            return existing, self._parse(suffix, content)

        text = self._parse(suffix, content)
        if not text.strip():
            hint = " (scanned PDFs require OCR and are not supported)" if suffix == ".pdf" else ""
            raise DocumentError(f"document contains no extractable text{hint}")
        object_key = f"projects/{project_id}/source/{digest}/{safe_name}"
        parsed_key = f"projects/{project_id}/parsed/{digest}/content.txt"
        document = Document(
            project_id=project_id,
            filename=safe_name,
            mime_type=mime_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream",
            size_bytes=len(content),
            content_hash=digest,
            object_key=object_key,
            parsed_object_key=parsed_key,
            status="pending",
        )
        self.repositories.documents.add(document)
        uploaded: list[str] = []
        try:
            self._put(object_key, content, document.mime_type)
            uploaded.append(object_key)
            self._put(parsed_key, text.encode("utf-8"), "text/plain; charset=utf-8")
            uploaded.append(parsed_key)
            chunks = self._chunks(project_id, document.id, text)
            for chunk in chunks:
                self.repositories.document_chunks.add(chunk)
            self._index_chunks(chunks, text)
            for chunk in chunks:
                self.repositories.document_chunks.update(
                    chunk.model_copy(update={"status": "ready"})
                )
            document = document.model_copy(update={"status": "ready"})
            self.repositories.documents.update(document)
            return document, text
        except Exception as exc:
            for key in uploaded:
                try:
                    self._delete(key)
                except Exception:
                    pass
            failed = document.model_copy(
                update={"status": "cleanup_required", "error_message": str(exc)[:1000]}
            )
            self.repositories.documents.update(failed)
            raise

    def _parse(self, suffix: str, content: bytes) -> str:
        if suffix in {".txt", ".md", ".markdown"}:
            for encoding in ("utf-8-sig", "utf-8", "gb18030"):
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise DocumentError("text encoding is not UTF-8 or GB18030")
        if suffix == ".docx":
            try:
                doc = DocxDocument(io.BytesIO(content))
                return "\n".join(paragraph.text for paragraph in doc.paragraphs)
            except Exception as exc:
                raise DocumentError("invalid DOCX document") from exc
        try:
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise DocumentError("invalid PDF document") from exc

    @staticmethod
    def _safe_filename(filename: str) -> str:
        name = Path(filename).name
        name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._ -]", "_", name).strip()
        if not name:
            raise DocumentError("filename is empty after sanitization")
        return name[:255]

    @staticmethod
    def _chunks(
        project_id: str, document_id: str, text: str, size: int = 1200, overlap: int = 120
    ) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        start = 0
        ordinal = 0
        while start < len(text):
            value = text[start : start + size]
            digest = hashlib.sha256(value.encode()).hexdigest()
            chunks.append(
                DocumentChunk(
                    project_id=project_id,
                    document_id=document_id,
                    ordinal=ordinal,
                    content_hash=digest,
                    preview=value[:240],
                    vector_id=f"{document_id}:{ordinal}",
                )
            )
            if start + size >= len(text):
                break
            start += size - overlap
            ordinal += 1
        return chunks

    def _index_chunks(self, chunks: list[DocumentChunk], text: str) -> None:
        if not self.vector_store or not self.embedding_provider:
            return
        sections = [text[index * 1080 : index * 1080 + 1200] for index in range(len(chunks))]
        vectors = self.embedding_provider.embed_documents(sections)
        records = [
            VectorRecord(
                id=chunk.vector_id or f"{chunk.document_id}:{chunk.ordinal}",
                project_id=chunk.project_id,
                source_id=chunk.document_id,
                source_type="document",
                chunk_ordinal=chunk.ordinal,
                content_hash=chunk.content_hash,
                metadata={"preview": chunk.preview},
                embedding=vector,
            )
            for chunk, vector in zip(chunks, vectors, strict=True)
        ]
        self.vector_store.upsert(records)

    def _put(self, key: str, data: bytes, content_type: str) -> None:
        method = getattr(self.object_store, "put_bytes", None) or self.object_store.put
        method(key, data, content_type=content_type)

    def _delete(self, key: str) -> None:
        method = getattr(self.object_store, "remove", None) or self.object_store.delete
        method(key)
